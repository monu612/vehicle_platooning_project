import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

md_file = "/Users/monu/vehicle_platooning_project/report_source.md"
output_docx = "/Users/monu/vehicle_platooning_project/Final_Submission_Report.docx"
template_docx = "/Users/monu/Desktop/vehicle_platooning_project/Format.docx"

# Read markdown
with open(md_file, "r") as f:
    md_text = f.read()

# Load document
doc = Document(template_docx)

def replace_in_doc(doc, search_text, replace_text):
    for p in doc.paragraphs:
        if search_text in p.text:
            for r in p.runs:
                if search_text in r.text:
                    r.text = r.text.replace(search_text, replace_text)
            if search_text in p.text:
                p.text = p.text.replace(search_text, replace_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if search_text in p.text:
                        for r in p.runs:
                            if search_text in r.text:
                                r.text = r.text.replace(search_text, replace_text)
                        if search_text in p.text:
                            p.text = p.text.replace(search_text, replace_text)

# Replace all template placeholders
replacements = {
    "TITLE OF THE PROJECT": "Adaptive Smart ACO Based Robust Communication for Vehicle Platooning using Spider-Web Topology",
    "Student Name (Roll Number)": "Monu (Final Year)",
    "Student Name: ": "Monu",
    "Roll Number: ": "",
    "<Supervisor Name>": "Project Supervisor",
    "<Designation>": "Assistant Professor",
    "Write your abstract here.": "This project addresses the critical challenge of reliable vehicle-to-vehicle (V2V) communication in next-generation autonomous vehicle platooning. Existing star topology communication architectures suffer from single points of failure, high latency, and significant performance degradation under network congestion or link failures. To overcome these limitations, we propose a novel decentralised communication architecture combining a Spider-Web Mesh Topology with Adaptive Smart Ant Colony Optimization (AS-ACO) for dynamic routing. The spider-web topology ensures high path redundancy and fault tolerance. Unlike classical ACO which uses fixed parameters, our Adaptive Smart ACO dynamically adjusts routing parameters (α, β, ρ) in real time based on network congestion, link reliability, and network instability. This adaptive intelligence continuously prevents network loops and ensures optimal convergence under high stress. Extensive simulations demonstrate that the proposed AS-ACO-driven spider-web architecture achieves near 100% packet delivery ratio, significantly reduces communication latency during unstable network conditions, and maintains robust connectivity even when up to 30% of communication links fail simultaneously. This report comprehensively details the mathematical models, algorithmic design, software architecture, and performance evaluations of the implemented system.",
    "Keywords: Write your keywords here separated by comma.": "Keywords: Vehicle Platooning, Adaptive Smart ACO, AS-ACO, Dynamic Parameter Optimization, Spider-Web Topology, V2V Communication, Fault Tolerance, Dynamic Routing, Autonomous Vehicles."
}

for search, replace in replacements.items():
    replace_in_doc(doc, search, replace)

# Build content from Markdown
sections_raw = re.split(r'\n# SECTION \d+ — [^\n]+', md_text)
titles = re.findall(r'\n# SECTION \d+ — ([^\n]+)', md_text)
sections = dict(zip(titles, sections_raw[1:]))

chapter_map = {
    "Chapter 1: Introduction": ["BIG PICTURE"],
    "Chapter 2: Literature Survey": ["DOMAIN KNOWLEDGE (FIRST PRINCIPLES)", "EXISTING SYSTEM ANALYSIS", "PROBLEM STATEMENT"],
    "Chapter 3: Proposed Methodology": ["PROPOSED SOLUTION", "SPIDER-WEB TOPOLOGY", "ANT COLONY OPTIMIZATION (VERY DEEP)", "MATHEMATICAL MODEL", "COMPLETE ALGORITHM DESIGN"],
    "Chapter 4: Implementation": ["SOFTWARE ARCHITECTURE", "COMPLETE CODE EXPLANATION", "SIMULATION DESIGN", "VISUALIZATION SYSTEM", "UI DESIGN"],
    "Chapter 5: Results and Discussion": ["TESTING SYSTEM", "PERFORMANCE METRICS", "BASELINE COMPARISON", "END-TO-END DATA FLOW"],
    "Chapter 6: Conclusion and Future Work": [],
    "Appendices": ["HOW TO BUILD FROM SCRATCH", "VIVA PREPARATION", "RESEARCH DEPTH"]
}

doc.add_page_break()

image_dir = "/Users/monu/vehicle_platooning_project/output"
chapter_index = 1

def add_md_to_doc(text, doc, chapter_num):
    lines = text.split('\n')
    in_code_block = False
    code_text = ""
    
    for line in lines:
        if line.startswith('```'):
            if in_code_block:
                # Add code block
                p = doc.add_paragraph(code_text)
                p.style.font.name = 'Courier New'
                p.style.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.5)
                in_code_block = False
                code_text = ""
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_text += line + "\n"
            continue

        if line.startswith('### '):
            doc.add_heading(line.replace('### ', '').strip(), level=3)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', '').strip(), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.replace('# ', '').strip(), level=2)
        elif line.strip().startswith('|') and '|---' in line:
            continue # simple skip of markdown table separator
        elif line.strip().startswith('|'):
            # simple representation of table rows as text
            doc.add_paragraph(line.strip().replace('|', '  '))
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            p = doc.add_paragraph('• ' + line.strip()[2:])
        elif re.match(r'^\d+\.\s', line.strip()):
            p = doc.add_paragraph(line.strip(),  )
        elif line.strip():
            # Bold parsing
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

for chapter_name, section_titles in chapter_map.items():
    doc.add_heading(chapter_name, level=1)
    for title in section_titles:
        if title in sections:
            doc.add_heading(title, level=2)
            add_md_to_doc(sections[title], doc, chapter_index)
            
            # Special image injection based on chapter content
            if "UI DESIGN" in title:
                img_path = os.path.join(image_dir, "latency_graph.png") # UI placeholder
                doc.add_paragraph("Figure: Vehicle Communication Dashboard UI")
            if "PERFORMANCE METRICS" in title:
                for img, caption in [
                    ("latency_graph.png", f"Figure {chapter_index}.1: Latency Comparison (ACO vs Baseline)"),
                    ("pdr_graph.png", f"Figure {chapter_index}.2: Packet Delivery Ratio"),
                    ("redundancy_graph.png", f"Figure {chapter_index}.3: Network Redundancy"),
                    ("convergence.png", f"Figure {chapter_index}.4: Pheromone Convergence"),
                    ("latency_distribution.png", f"Figure {chapter_index}.5: Latency Distribution"),
                    ("sensitivity.png", f"Figure {chapter_index}.6: ACO Parameter Sensitivity")
                ]:
                    img_p = os.path.join(image_dir, img)
                    if os.path.exists(img_p):
                        doc.add_picture(img_p, width=Inches(5.5))
                        p = doc.add_paragraph(caption)
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    chapter_index += 1

doc.save(output_docx)
print("Saved DOCX to:", output_docx)
